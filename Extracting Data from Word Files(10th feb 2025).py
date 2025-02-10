#!/usr/bin/env python
# coding: utf-8

# In[1]:


get_ipython().system('pip install python-docx')


# In[2]:


import docx


# In[6]:


doc=open(r"C:\Users\rajuc\Downloads\039.docx","rb")
document=docx.Document(doc)


# In[7]:


docu=""
for para in document.paragraphs:
    docu +=para.text
print(docu)


# In[12]:


for i in range(len(document.paragraphs)):
    print("The content of the paragraph"+str(i)+"is:"+document.paragraphs[i].text+"\n")


# In[ ]:




